import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ============ STYLES ============
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.bold = True
    if i == 1:
        h.font.size = Pt(16)
    elif i == 2:
        h.font.size = Pt(14)
    else:
        h.font.size = Pt(12)

# ============ COVER PAGE ============
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('LAPORAN PRAKTIKUM\nPEMROGRAMAN BERORIENTASI OBJEK')
r.bold = True
r.font.size = Pt(18)
r.font.name = 'Times New Roman'

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SISTEM INFORMASI PERPUSTAKAAN\n(Library Management System)')
r.bold = True
r.font.size = Pt(16)
r.font.name = 'Times New Roman'

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Java Swing + JDBC + MVC Architecture')
r.font.size = Pt(14)
r.font.name = 'Times New Roman'
r.italic = True

for _ in range(4):
    doc.add_paragraph('')

info_lines = [
    'Disusun Oleh:',
    'Nama    : [Nama Mahasiswa]',
    'NIM     : [NIM]',
    'Kelas   : [Kelas]',
    'Dosen   : [Nama Dosen]',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

for _ in range(3):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PROGRAM STUDI TEKNIK INFORMATIKA\nFAKULTAS TEKNIK\nTAHUN AJARAN 2025/2026')
r.font.size = Pt(12)
r.font.name = 'Times New Roman'
r.bold = True

doc.add_page_break()

# ============ DAFTAR ISI ============
doc.add_heading('DAFTAR ISI', level=1)
toc_items = [
    ('BAB I - PENDAHULUAN', '1'),
    ('  1.1 Latar Belakang', '1'),
    ('  1.2 Tujuan Praktikum', '1'),
    ('  1.3 Ruang Lingkup', '2'),
    ('BAB II - DASAR TEORI', '3'),
    ('  2.1 Object-Oriented Programming (OOP)', '3'),
    ('  2.2 Java Swing', '3'),
    ('  2.3 JDBC (Java Database Connectivity)', '3'),
    ('  2.4 Pola Arsitektur MVC', '3'),
    ('  2.5 Design Pattern: DAO', '4'),
    ('BAB III - PERANCANGAN SISTEM', '5'),
    ('  3.1 Arsitektur Aplikasi', '5'),
    ('  3.2 Struktur Proyek', '5'),
    ('  3.3 Desain Database', '6'),
    ('BAB IV - IMPLEMENTASI (SOURCE CODE)', '7'),
    ('  4.1 Database Setup (SQL)', '7'),
    ('  4.2 Layer Model', '7'),
    ('  4.3 Layer DAO (Data Access Object)', '7'),
    ('  4.4 Layer Controller', '7'),
    ('  4.5 Layer View (GUI)', '7'),
    ('  4.6 Exception Handling', '7'),
    ('  4.7 Utility Classes', '7'),
    ('BAB V - FITUR APLIKASI', '8'),
    ('BAB VI - KESIMPULAN', '9'),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item)
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

doc.add_page_break()

# ============ BAB I ============
doc.add_heading('BAB I - PENDAHULUAN', level=1)
doc.add_heading('1.1 Latar Belakang', level=2)
doc.add_paragraph(
    'Perpustakaan merupakan salah satu fasilitas penting dalam dunia pendidikan yang memerlukan '
    'sistem pengelolaan yang efisien. Pengelolaan data buku, anggota, dan transaksi peminjaman '
    'secara manual seringkali menimbulkan permasalahan seperti data yang tidak akurat, proses '
    'yang lambat, dan kesulitan dalam pencarian informasi. Oleh karena itu, diperlukan sebuah '
    'sistem informasi berbasis komputer yang dapat mengelola seluruh aktivitas perpustakaan '
    'secara digital.')
doc.add_paragraph(
    'Pada praktikum ini, dikembangkan sebuah aplikasi Sistem Informasi Perpustakaan (Library '
    'Management System) menggunakan bahasa pemrograman Java dengan pendekatan Object-Oriented '
    'Programming (OOP). Aplikasi ini dibangun menggunakan Java Swing untuk antarmuka grafis (GUI), '
    'JDBC untuk konektivitas database MySQL, dan mengikuti pola arsitektur Model-View-Controller (MVC) '
    'serta pattern Data Access Object (DAO) untuk pemisahan tanggung jawab yang baik.')

doc.add_heading('1.2 Tujuan Praktikum', level=2)
tujuan = [
    'Mengimplementasikan konsep Object-Oriented Programming (OOP) meliputi: Inheritance, Polymorphism, Abstraction, Encapsulation, dan Interface.',
    'Membangun aplikasi GUI desktop menggunakan Java Swing.',
    'Mengimplementasikan koneksi database MySQL menggunakan JDBC dengan PreparedStatement.',
    'Menerapkan arsitektur MVC (Model-View-Controller) dan pattern DAO.',
    'Mengimplementasikan Exception Handling menggunakan Checked dan Unchecked Exception.',
    'Memanfaatkan Java Collections Framework (ArrayList dan HashMap) untuk caching data.',
    'Mengimplementasikan fitur File I/O untuk export CSV, logging, dan cetak struk.',
]
for t in tujuan:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('1.3 Ruang Lingkup', level=2)
doc.add_paragraph(
    'Aplikasi ini mencakup modul manajemen buku (CRUD, pencarian, export CSV), manajemen anggota '
    '(CRUD, pencarian, validasi email), transaksi peminjaman dan pengembalian (dengan validasi stok otomatis), '
    'dashboard statistik, sistem login sederhana, perhitungan denda keterlambatan, dan cetak struk peminjaman.')

doc.add_page_break()

# ============ BAB II ============
doc.add_heading('BAB II - DASAR TEORI', level=1)
doc.add_heading('2.1 Object-Oriented Programming (OOP)', level=2)
doc.add_paragraph(
    'OOP adalah paradigma pemrograman yang mengorganisasi perangkat lunak ke dalam objek-objek '
    'yang memiliki atribut (data) dan method (perilaku). Empat pilar utama OOP yang diterapkan: '
    'Encapsulation (pembungkusan data dengan akses modifier private dan getter/setter), '
    'Inheritance (pewarisan sifat dari class induk BaseEntity), '
    'Polymorphism (override method getId() dan getDisplayInfo() di setiap subclass), dan '
    'Abstraction (abstract class BaseEntity dan interface Exportable).')

doc.add_heading('2.2 Java Swing', level=2)
doc.add_paragraph(
    'Java Swing adalah library GUI yang menyediakan komponen-komponen untuk membangun '
    'antarmuka pengguna desktop, seperti JFrame, JPanel, JTable, JTabbedPane, JComboBox, '
    'JTextField, JButton, dan JOptionPane.')

doc.add_heading('2.3 JDBC (Java Database Connectivity)', level=2)
doc.add_paragraph(
    'JDBC adalah API Java untuk mengakses database relasional. Pada proyek ini digunakan '
    'MySQL Connector/J sebagai JDBC driver, PreparedStatement untuk mencegah SQL Injection, '
    'dan try-with-resources untuk pengelolaan koneksi yang aman.')

doc.add_heading('2.4 Pola Arsitektur MVC', level=2)
doc.add_paragraph(
    'MVC membagi aplikasi menjadi tiga komponen: Model (representasi data/entitas), '
    'View (antarmuka pengguna/GUI), dan Controller (logika bisnis yang menjembatani Model dan View). '
    'Pemisahan ini memudahkan maintenance dan pengembangan.')

doc.add_heading('2.5 Design Pattern: DAO (Data Access Object)', level=2)
doc.add_paragraph(
    'DAO memisahkan logika akses data (query SQL) dari logika bisnis. '
    'Setiap entitas memiliki kelas DAO tersendiri (BukuDAO, AnggotaDAO, PeminjamanDAO) '
    'yang menangani operasi CRUD ke database.')

doc.add_page_break()

# ============ BAB III ============
doc.add_heading('BAB III - PERANCANGAN SISTEM', level=1)
doc.add_heading('3.1 Arsitektur Aplikasi', level=2)
doc.add_paragraph(
    'Aplikasi menggunakan arsitektur berlapis (layered architecture) dengan alur: '
    'View (GUI/Swing) → Controller (Business Logic + Cache) → DAO (Data Access) → Database (MySQL). '
    'Setiap layer memiliki tanggung jawab yang terpisah sehingga memudahkan pengembangan dan debugging.')

doc.add_heading('3.2 Struktur Proyek', level=2)
struktur = """LibraryManagementSystem/
├── database_setup.sql
├── src/
│   ├── model/
│   │   ├── BaseEntity.java       (Abstract Class)
│   │   ├── Exportable.java       (Interface)
│   │   ├── Buku.java
│   │   ├── Anggota.java
│   │   └── Peminjaman.java
│   ├── view/
│   │   ├── LoginFrame.java
│   │   ├── MainFrame.java
│   │   ├── DashboardPanel.java
│   │   ├── BukuPanel.java
│   │   ├── AnggotaPanel.java
│   │   └── PeminjamanPanel.java
│   ├── controller/
│   │   ├── BukuController.java
│   │   ├── AnggotaController.java
│   │   └── PeminjamanController.java
│   ├── dao/
│   │   ├── BukuDAO.java
│   │   ├── AnggotaDAO.java
│   │   └── PeminjamanDAO.java
│   ├── exception/
│   │   ├── BukuTidakTersediaException.java
│   │   ├── StokHabisException.java
│   │   └── DataTidakValidException.java
│   └── util/
│       ├── DBConnection.java
│       ├── FileExporter.java
│       └── LoggerUtil.java
└── lib/   ← mysql-connector-j-x.x.x.jar"""

p = doc.add_paragraph()
r = p.add_run(struktur)
r.font.name = 'Consolas'
r.font.size = Pt(9)

doc.add_heading('3.3 Desain Database', level=2)
doc.add_paragraph('Database perpustakaan_db terdiri dari 3 tabel utama:')

# Table: buku
doc.add_paragraph('Tabel buku:', style='List Bullet')
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Kolom', 'Tipe Data', 'Keterangan']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
buku_cols = [
    ('kode_buku', 'VARCHAR(10)', 'PRIMARY KEY'),
    ('judul', 'VARCHAR(200)', 'NOT NULL'),
    ('pengarang', 'VARCHAR(100)', 'NOT NULL'),
    ('penerbit', 'VARCHAR(100)', ''),
    ('tahun_terbit', 'INT', ''),
    ('stok', 'INT', 'DEFAULT 1'),
    ('created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP'),
]
for i, (col, tipe, ket) in enumerate(buku_cols, 1):
    table.rows[i].cells[0].text = col
    table.rows[i].cells[1].text = tipe
    table.rows[i].cells[2].text = ket

doc.add_paragraph('')
doc.add_paragraph('Tabel anggota:', style='List Bullet')
table2 = doc.add_table(rows=7, cols=3)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
anggota_cols = [
    ('id_anggota', 'VARCHAR(10)', 'PRIMARY KEY'),
    ('nama', 'VARCHAR(100)', 'NOT NULL'),
    ('alamat', 'TEXT', ''),
    ('no_telp', 'VARCHAR(15)', ''),
    ('email', 'VARCHAR(50)', ''),
    ('tgl_daftar', 'DATE', 'DEFAULT CURRENT_DATE'),
]
for i, (col, tipe, ket) in enumerate(anggota_cols, 1):
    table2.rows[i].cells[0].text = col
    table2.rows[i].cells[1].text = tipe
    table2.rows[i].cells[2].text = ket

doc.add_paragraph('')
doc.add_paragraph('Tabel peminjaman:', style='List Bullet')
table3 = doc.add_table(rows=7, cols=3)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers):
    table3.rows[0].cells[i].text = h
    for p in table3.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
pinjam_cols = [
    ('id_pinjam', 'INT', 'PRIMARY KEY, AUTO_INCREMENT'),
    ('kode_buku', 'VARCHAR(10)', 'FOREIGN KEY → buku'),
    ('id_anggota', 'VARCHAR(10)', 'FOREIGN KEY → anggota'),
    ('tgl_pinjam', 'DATE', 'NOT NULL'),
    ('tgl_kembali', 'DATE', 'NULL'),
    ('status', "ENUM('DIPINJAM','KEMBALI')", "DEFAULT 'DIPINJAM'"),
]
for i, (col, tipe, ket) in enumerate(pinjam_cols, 1):
    table3.rows[i].cells[0].text = col
    table3.rows[i].cells[1].text = tipe
    table3.rows[i].cells[2].text = ket

doc.add_page_break()

# ============ BAB IV - SOURCE CODE ============
doc.add_heading('BAB IV - IMPLEMENTASI (SOURCE CODE)', level=1)

def add_code(doc, title, filepath, heading_level=3):
    """Read a file and add it as formatted code block."""
    doc.add_heading(title, level=heading_level)
    base = os.path.basename(filepath)
    doc.add_paragraph(f'File: {filepath}')
    
    with open(filepath, 'r', encoding='utf-8') as f:
        code = f.read()
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(code)
    r.font.name = 'Consolas'
    r.font.size = Pt(8)

BASE = r'c:\Users\Administrator\Downloads\LibraryManagementSystem'

# 4.1 Database
doc.add_heading('4.1 Database Setup (SQL)', level=2)
add_code(doc, 'database_setup.sql', os.path.join(BASE, 'database_setup.sql'))

# 4.2 Model
doc.add_heading('4.2 Layer Model', level=2)
doc.add_paragraph('Layer Model berisi class-class yang merepresentasikan entitas data. '
    'Menggunakan abstract class BaseEntity sebagai parent dan interface Exportable untuk fitur export.')
add_code(doc, 'BaseEntity.java (Abstract Class)', os.path.join(BASE, 'src', 'model', 'BaseEntity.java'))
add_code(doc, 'Exportable.java (Interface)', os.path.join(BASE, 'src', 'model', 'Exportable.java'))
add_code(doc, 'Buku.java', os.path.join(BASE, 'src', 'model', 'Buku.java'))
add_code(doc, 'Anggota.java', os.path.join(BASE, 'src', 'model', 'Anggota.java'))
add_code(doc, 'Peminjaman.java', os.path.join(BASE, 'src', 'model', 'Peminjaman.java'))

# 4.3 DAO
doc.add_heading('4.3 Layer DAO (Data Access Object)', level=2)
doc.add_paragraph('Layer DAO menangani seluruh operasi CRUD ke database menggunakan PreparedStatement dan try-with-resources.')
add_code(doc, 'BukuDAO.java', os.path.join(BASE, 'src', 'dao', 'BukuDAO.java'))
add_code(doc, 'AnggotaDAO.java', os.path.join(BASE, 'src', 'dao', 'AnggotaDAO.java'))
add_code(doc, 'PeminjamanDAO.java', os.path.join(BASE, 'src', 'dao', 'PeminjamanDAO.java'))

# 4.4 Controller
doc.add_heading('4.4 Layer Controller', level=2)
doc.add_paragraph('Layer Controller berisi logika bisnis, validasi data, dan caching menggunakan ArrayList & HashMap.')
add_code(doc, 'BukuController.java', os.path.join(BASE, 'src', 'controller', 'BukuController.java'))
add_code(doc, 'AnggotaController.java', os.path.join(BASE, 'src', 'controller', 'AnggotaController.java'))
add_code(doc, 'PeminjamanController.java', os.path.join(BASE, 'src', 'controller', 'PeminjamanController.java'))

# 4.5 View
doc.add_heading('4.5 Layer View (GUI)', level=2)
doc.add_paragraph('Layer View menggunakan Java Swing untuk membangun antarmuka pengguna berbasis GUI desktop.')
add_code(doc, 'LoginFrame.java', os.path.join(BASE, 'src', 'view', 'LoginFrame.java'))
add_code(doc, 'MainFrame.java', os.path.join(BASE, 'src', 'view', 'MainFrame.java'))
add_code(doc, 'DashboardPanel.java', os.path.join(BASE, 'src', 'view', 'DashboardPanel.java'))
add_code(doc, 'BukuPanel.java', os.path.join(BASE, 'src', 'view', 'BukuPanel.java'))
add_code(doc, 'AnggotaPanel.java', os.path.join(BASE, 'src', 'view', 'AnggotaPanel.java'))
add_code(doc, 'PeminjamanPanel.java', os.path.join(BASE, 'src', 'view', 'PeminjamanPanel.java'))

# 4.6 Exception
doc.add_heading('4.6 Custom Exception Handling', level=2)
doc.add_paragraph('Aplikasi menggunakan 3 custom exception: 1 checked exception (BukuTidakTersediaException) '
    'dan 2 unchecked exception (StokHabisException, DataTidakValidException).')
add_code(doc, 'BukuTidakTersediaException.java (Checked)', os.path.join(BASE, 'src', 'exception', 'BukuTidakTersediaException.java'))
add_code(doc, 'StokHabisException.java (Unchecked)', os.path.join(BASE, 'src', 'exception', 'StokHabisException.java'))
add_code(doc, 'DataTidakValidException.java (Unchecked)', os.path.join(BASE, 'src', 'exception', 'DataTidakValidException.java'))

# 4.7 Utility
doc.add_heading('4.7 Utility Classes', level=2)
doc.add_paragraph('Utility classes menyediakan fungsi pembantu: koneksi database, export file CSV/struk, dan logging.')
add_code(doc, 'DBConnection.java', os.path.join(BASE, 'src', 'util', 'DBConnection.java'))
add_code(doc, 'FileExporter.java', os.path.join(BASE, 'src', 'util', 'FileExporter.java'))
add_code(doc, 'LoggerUtil.java', os.path.join(BASE, 'src', 'util', 'LoggerUtil.java'))

doc.add_page_break()

# ============ BAB V ============
doc.add_heading('BAB V - FITUR APLIKASI', level=1)
doc.add_paragraph('Berikut adalah fitur-fitur yang telah berhasil diimplementasikan:')

fitur_list = [
    ('Login System', 'Autentikasi petugas perpustakaan sebelum mengakses sistem (default: admin/admin123).'),
    ('Dashboard Statistik', 'Menampilkan total buku, total anggota, dan jumlah buku yang sedang dipinjam secara real-time.'),
    ('CRUD Buku', 'Tambah, edit, hapus, dan lihat data buku dengan validasi kode unik dan stok non-negatif.'),
    ('Pencarian Buku', 'Pencarian berdasarkan judul atau nama pengarang.'),
    ('Export Buku ke CSV', 'Mengekspor seluruh data buku ke file CSV menggunakan JFileChooser.'),
    ('CRUD Anggota', 'Tambah, edit, hapus, dan lihat data anggota dengan validasi ID unik dan format email.'),
    ('Pencarian Anggota', 'Pencarian berdasarkan nama atau ID anggota.'),
    ('Transaksi Peminjaman', 'Proses peminjaman buku oleh anggota dengan validasi ketersediaan stok secara otomatis.'),
    ('Transaksi Pengembalian', 'Proses pengembalian buku dengan update status dan pengembalian stok otomatis.'),
    ('Perhitungan Denda', 'Denda keterlambatan dihitung otomatis: Rp 1.000/hari jika melebihi 7 hari masa pinjam.'),
    ('Cetak Struk Peminjaman', 'Mencetak struk peminjaman ke file teks (.txt) dengan informasi lengkap transaksi.'),
    ('Riwayat Transaksi', 'Menampilkan seluruh riwayat peminjaman dan pengembalian.'),
    ('Logging Aktivitas', 'Setiap transaksi dicatat di file log.txt dengan timestamp otomatis.'),
]
for title, desc in fitur_list:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(title + ': ')
    r.bold = True
    p.add_run(desc)

doc.add_heading('Konsep OOP yang Diterapkan', level=2)
oop_list = [
    ('Encapsulation', 'Atribut private dengan getter/setter di semua model class.'),
    ('Inheritance', 'Buku, Anggota, Peminjaman meng-extend abstract class BaseEntity.'),
    ('Polymorphism', 'Override method getId() dan getDisplayInfo() di setiap subclass.'),
    ('Abstraction', 'Abstract class BaseEntity dan interface Exportable.'),
    ('Interface', 'Exportable interface diimplementasikan oleh Buku dan Anggota untuk fitur export.'),
]
for title, desc in oop_list:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(title + ': ')
    r.bold = True
    p.add_run(desc)

doc.add_heading('Cara Menjalankan Aplikasi', level=2)
steps = [
    'Jalankan XAMPP dan aktifkan MySQL.',
    'Buka phpMyAdmin, jalankan file database_setup.sql untuk membuat database perpustakaan_db.',
    'Download MySQL Connector/J dan letakkan file .jar di folder lib/.',
    'Sesuaikan konfigurasi koneksi di DBConnection.java (user dan password MySQL).',
    'Compile: javac -cp "lib/*" -d bin src/model/*.java src/exception/*.java src/util/*.java src/dao/*.java src/controller/*.java src/view/*.java',
    'Jalankan: java -cp "bin;lib/*" view.MainFrame',
    'Login dengan username: admin, password: admin123.',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# ============ BAB VI ============
doc.add_heading('BAB VI - KESIMPULAN', level=1)
doc.add_heading('6.1 Kesimpulan', level=2)
doc.add_paragraph(
    'Dari praktikum ini dapat disimpulkan bahwa:')
kesimpulan = [
    'Konsep OOP (Encapsulation, Inheritance, Polymorphism, Abstraction) berhasil diterapkan secara menyeluruh dalam pengembangan Sistem Informasi Perpustakaan.',
    'Arsitektur MVC dan pattern DAO membantu memisahkan tanggung jawab setiap komponen sehingga kode lebih terstruktur, mudah dipahami, dan mudah di-maintain.',
    'Java Swing dapat digunakan untuk membangun aplikasi GUI desktop yang interaktif dan user-friendly.',
    'JDBC dengan PreparedStatement dan try-with-resources menyediakan cara yang aman dan efisien untuk mengakses database MySQL.',
    'Custom Exception Handling (checked dan unchecked) membantu menangani error secara spesifik dan memberikan pesan yang informatif kepada pengguna.',
    'Java Collections Framework (ArrayList dan HashMap) efektif digunakan untuk caching data di memory sehingga mempercepat pencarian.',
    'Fitur File I/O (export CSV, logging, cetak struk) berhasil diimplementasikan untuk mendukung kebutuhan operasional perpustakaan.',
]
for k in kesimpulan:
    doc.add_paragraph(k, style='List Bullet')

doc.add_heading('6.2 Saran', level=2)
saran = [
    'Mengimplementasikan hashing password untuk keamanan sistem login.',
    'Menambahkan fitur laporan statistik yang lebih detail (grafik/chart).',
    'Mengembangkan fitur multi-user dengan role-based access control.',
    'Migrasi ke framework yang lebih modern seperti JavaFX untuk tampilan yang lebih menarik.',
    'Menambahkan fitur backup dan restore database.',
]
for s in saran:
    doc.add_paragraph(s, style='List Bullet')

# ============ SAVE ============
output_path = os.path.join(BASE, 'Laporan_Praktikum_Library_Management_System.docx')
doc.save(output_path)
print(f'Laporan berhasil dibuat: {output_path}')
